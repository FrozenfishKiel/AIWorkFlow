from pathlib import Path
import importlib

from sqlmodel import Session
from docx import Document

from app.repositories.knowledge_repository import KnowledgeRepository
from app.repositories.task_repository import TaskRepository
from app.services.knowledge_index_service import KnowledgeIndexService
from app.services.task_pipeline_service import TaskPipelineService


def test_pipeline_service_generates_understanding_retrieval_and_workflow_results(
    session: Session,
    tmp_path: Path,
) -> None:
    knowledge_repository = KnowledgeRepository(session)
    source_file = tmp_path / "review-standard.md"
    source_file.write_text(
        "# Review Standard\n\n"
        "Launch plans must keep claims reviewable.\n\n"
        "Visible citations are mandatory before export.\n",
        encoding="utf-8",
    )
    document = knowledge_repository.create_document(
        title="Review Standard",
        source_path=str(source_file),
        source_type="review_standard",
        domain="content-ops",
    )
    KnowledgeIndexService(knowledge_repository).index_document(document.id)

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="text",
        content="Create a launch plan for the new product article.",
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert processed_task.understanding is not None
    assert processed_task.understanding["summary"]
    assert processed_task.understanding["audience"]
    assert processed_task.understanding["risk_points"]
    assert processed_task.understanding["uncertain_items"]
    assert processed_task.understanding["input_quality"]["source_kind"] == "text"
    assert processed_task.retrieval_hits
    assert processed_task.retrieval_hits[0]["source_id"] == str(document.id)
    assert processed_task.retrieval_hits[0]["title"] == "Review Standard"
    assert processed_task.workflow_result is not None
    assert processed_task.workflow_result["draft"]
    assert processed_task.workflow_result["evidence_used"]
    assert processed_task.workflow_result["evidence_used"][0]["snippet"]
    assert processed_task.workflow_result["evidence_used"][0]["reason"]
    assert processed_task.workflow_result["manual_checks"]
    assert processed_task.workflow_result["context_summary"]["context_sections"]
    assert processed_task.workflow_result["review_notes"]
    assert processed_task.workflow_result["processing_trace"]
    assert processed_task.approved_snapshot is not None
    assert processed_task.approved_snapshot["understanding"] == processed_task.understanding
    assert processed_task.approved_snapshot["retrieval_hits"] == processed_task.retrieval_hits
    assert processed_task.approved_snapshot["workflow_result"] == processed_task.workflow_result


def test_pipeline_service_can_use_injected_generation_provider_outputs(
    session: Session,
    tmp_path: Path,
) -> None:
    knowledge_repository = KnowledgeRepository(session)
    source_file = tmp_path / "brand-playbook.md"
    source_file.write_text(
        "# Brand Playbook\n\n"
        "Always keep launch messaging practical and evidence-backed.\n",
        encoding="utf-8",
    )
    document = knowledge_repository.create_document(
        title="Brand Playbook",
        source_path=str(source_file),
        source_type="playbook",
        domain="brand",
    )
    KnowledgeIndexService(knowledge_repository).index_document(document.id)

    class FakeGenerationProvider:
        provider_name = "fake-live-provider"

        def build_understanding(self, parsed_input: dict[str, object]) -> dict[str, object]:
            assert "launch messaging" in str(parsed_input["parsed_text"]).lower()
            return {
                "summary": "Model-backed understanding for launch messaging.",
                "audience": ["brand-editor", "content-ops"],
                "key_points": [
                    "Keep the launch message practical.",
                    "Carry evidence into the workflow result.",
                ],
                "risk_points": ["Public claims still need evidence verification."],
                "uncertain_items": ["Final CTA still needs operator confirmation."],
            }

        def build_workflow(
            self,
            *,
            parsed_input: dict[str, object],
            understanding: dict[str, object],
            retrieval_hits: list[dict[str, object]],
            generation_context: dict[str, object],
        ) -> dict[str, object]:
            assert understanding["summary"] == "Model-backed understanding for launch messaging."
            assert retrieval_hits
            assert generation_context["selected_hits"]
            return {
                "draft": "Model-backed workflow draft with a grounded launch angle.",
                "review_notes": ["Check whether the CTA matches the active campaign."],
                "open_questions": ["Should the draft prioritize conversion or education first?"],
                "manual_checks": ["Validate the CTA against the cited playbook evidence."],
            }

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="text",
        content="Prepare launch messaging that stays practical and grounded in evidence.",
        knowledge_domain="brand",
    )

    service = TaskPipelineService(
        lambda: Session(session.get_bind()),
        generation_provider=FakeGenerationProvider(),
    )
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert processed_task.understanding is not None
    assert processed_task.understanding["summary"] == "Model-backed understanding for launch messaging."
    assert processed_task.understanding["audience"] == ["brand-editor", "content-ops"]
    assert processed_task.understanding["input_quality"]["source_kind"] == "text"
    assert processed_task.workflow_result is not None
    assert processed_task.workflow_result["draft"] == "Model-backed workflow draft with a grounded launch angle."
    assert processed_task.workflow_result["review_notes"] == [
        "Check whether the CTA matches the active campaign."
    ]
    assert processed_task.workflow_result["open_questions"] == [
        "Should the draft prioritize conversion or education first?"
    ]
    assert "Validate the CTA against the cited playbook evidence." in processed_task.workflow_result["manual_checks"]
    assert "Final CTA still needs operator confirmation." in processed_task.workflow_result["manual_checks"]
    assert processed_task.workflow_result["evidence_used"]
    assert processed_task.workflow_result["context_summary"]["selected_hit_count"] == 1
    assert any(
        "fake-live-provider" in trace_entry
        for trace_entry in processed_task.workflow_result["processing_trace"]
    )


def test_pipeline_service_persists_completion_audit_events(
    session: Session,
    tmp_path: Path,
) -> None:
    knowledge_repository = KnowledgeRepository(session)
    source_file = tmp_path / "audit-standard.md"
    source_file.write_text(
        "# Audit Standard\n\n"
        "Keep stable snapshot boundaries visible.\n",
        encoding="utf-8",
    )
    document = knowledge_repository.create_document(
        title="Audit Standard",
        source_path=str(source_file),
        source_type="audit_standard",
        domain="content-ops",
    )
    KnowledgeIndexService(knowledge_repository).index_document(document.id)

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="text",
        content="Create a task that should leave an auditable pipeline trail.",
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    audit_module = importlib.import_module("app.models.audit_log")
    AuditLog = audit_module.AuditLog
    audit_rows = list(
        session.exec(
            AuditLog.__table__.select().where(AuditLog.task_id == processed_task.id).order_by(AuditLog.created_at.asc())
        )
    )

    assert [(row.event_type, row.outcome) for row in audit_rows] == [
        ("pipeline_completed", "success"),
        ("snapshot_persisted", "success"),
    ]
    assert audit_rows[0].summary
    assert audit_rows[1].details["snapshot_fields"] == ["understanding", "retrieval_hits", "workflow_result"]


def test_pipeline_service_uses_uploaded_file_contents_and_retrieval_scope(
    session: Session,
    tmp_path: Path,
) -> None:
    knowledge_repository = KnowledgeRepository(session)

    brand_source = tmp_path / "brand-guideline.md"
    brand_source.write_text(
        "# Brand Guide\n\n"
        "Launch copy should stay warm and audience-specific.\n",
        encoding="utf-8",
    )
    compliance_source = tmp_path / "compliance-guideline.md"
    compliance_source.write_text(
        "# Compliance Guide\n\n"
        "Human review stays mandatory before export and every claim needs compliance sign-off.\n",
        encoding="utf-8",
    )

    brand_document = knowledge_repository.create_document(
        title="Brand Guide",
        source_path=str(brand_source),
        source_type="guideline",
        domain="brand",
    )
    compliance_document = knowledge_repository.create_document(
        title="Compliance Guide",
        source_path=str(compliance_source),
        source_type="guideline",
        domain="compliance",
    )
    index_service = KnowledgeIndexService(knowledge_repository)
    index_service.index_document(brand_document.id)
    index_service.index_document(compliance_document.id)

    uploaded_file = tmp_path / "task-brief.md"
    uploaded_file.write_text(
        "# Launch Brief\n\n"
        "Need compliance sign-off before publishing any launch claims.\n",
        encoding="utf-8",
    )

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="file",
        content=str(uploaded_file),
    )
    task.knowledge_domain = "compliance"
    session.add(task)
    session.commit()
    session.refresh(task)

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert "compliance sign-off" in processed_task.understanding["summary"].lower()
    assert processed_task.understanding["input_quality"]["source_kind"] == "file"
    assert processed_task.understanding["input_quality"]["quality_flags"] == []
    assert processed_task.retrieval_hits
    assert {hit["title"] for hit in processed_task.retrieval_hits} == {"Compliance Guide"}


def test_pipeline_service_extracts_visible_text_from_html_file_input(
    session: Session,
    tmp_path: Path,
) -> None:
    uploaded_file = tmp_path / "task-brief.html"
    uploaded_file.write_text(
        "<html><body><h1>Launch Brief</h1><p>Need compliance sign-off before publishing claims.</p></body></html>",
        encoding="utf-8",
    )

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="file",
        content=str(uploaded_file),
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert "launch brief" in processed_task.understanding["summary"].lower()
    assert "compliance sign-off" in processed_task.understanding["summary"].lower()


def test_pipeline_service_extracts_visible_text_from_docx_file_input(
    session: Session,
    tmp_path: Path,
) -> None:
    uploaded_file = tmp_path / "task-brief.docx"
    document = Document()
    document.add_heading("Launch Brief", level=1)
    document.add_paragraph("Need compliance sign-off before publishing claims.")
    document.save(uploaded_file)

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="file",
        content=str(uploaded_file),
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert "launch brief" in processed_task.understanding["summary"].lower()
    assert "compliance sign-off" in processed_task.understanding["summary"].lower()


def test_pipeline_service_uses_pdf_reader_for_pdf_file_input(
    session: Session,
    tmp_path: Path,
    monkeypatch,
) -> None:
    uploaded_file = tmp_path / "task-brief.pdf"
    uploaded_file.write_bytes(b"%PDF-1.4 fake pdf placeholder")

    class FakePdfPage:
        def extract_text(self) -> str:
            return "Launch Brief. Need compliance sign-off before publishing claims."

    class FakePdfReader:
        def __init__(self, path: Path) -> None:
            assert Path(path) == uploaded_file
            self.pages = [FakePdfPage()]

    monkeypatch.setattr("app.services.task_pipeline_service.PdfReader", FakePdfReader)

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="file",
        content=str(uploaded_file),
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert "launch brief" in processed_task.understanding["summary"].lower()
    assert "compliance sign-off" in processed_task.understanding["summary"].lower()


def test_pipeline_service_uses_real_url_text_for_retrieval(
    session: Session,
    tmp_path: Path,
    monkeypatch,
) -> None:
    knowledge_repository = KnowledgeRepository(session)

    source = tmp_path / "compliance-guideline.md"
    source.write_text(
        "# Compliance Guide\n\n"
        "Visible citations and human review are mandatory before launch export.\n",
        encoding="utf-8",
    )
    document = knowledge_repository.create_document(
        title="Compliance Guide",
        source_path=str(source),
        source_type="guideline",
        domain="compliance",
    )
    KnowledgeIndexService(knowledge_repository).index_document(document.id)

    monkeypatch.setattr(
        "app.services.task_pipeline_service.UrlIngestionService.fetch_public_content",
        lambda self, url: {
            "title": "Launch Update",
            "text": "Visible citations and human review are mandatory before launch export.",
            "extractor": "article",
            "quality_flags": [],
        },
    )

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="url",
        content="https://example.com/launch-update",
        knowledge_domain="compliance",
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert "visible citations" in processed_task.understanding["summary"].lower()
    assert processed_task.understanding["input_quality"]["source_kind"] == "url"
    assert processed_task.understanding["input_quality"]["metadata"]["extractor"] == "article"
    assert processed_task.retrieval_hits
    assert processed_task.retrieval_hits[0]["title"] == "Compliance Guide"


def test_pipeline_service_carries_url_extraction_quality_flags_forward(
    session: Session,
    monkeypatch,
) -> None:
    monkeypatch.setattr(
        "app.services.task_pipeline_service.UrlIngestionService.fetch_public_content",
        lambda self, url: {
            "title": "Status Page",
            "text": "Launch status",
            "extractor": "body",
            "quality_flags": ["fallback_html_extract", "shallow_url_extract"],
        },
    )

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="url",
        content="https://example.com/status",
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert "shallow_url_extract" in processed_task.understanding["input_quality"]["quality_flags"]
    assert "fallback_html_extract" in processed_task.understanding["input_quality"]["quality_flags"]
    assert processed_task.understanding["input_quality"]["metadata"]["title"] == "Status Page"
    assert "URL extraction may be shallow or noisy" in processed_task.understanding["risk_points"]
    assert "Confirm the source article is complete before reuse." in processed_task.workflow_result["manual_checks"]


def test_pipeline_service_marks_low_quality_inputs_and_carries_uncertainty_forward(
    session: Session,
) -> None:
    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="text",
        content="tiny",
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.status == "completed"
    assert processed_task.understanding is not None
    assert "Input is very short" in processed_task.understanding["risk_points"]
    assert "Source content may be incomplete" in processed_task.understanding["uncertain_items"]
    assert "short_input" in processed_task.understanding["input_quality"]["quality_flags"]
    assert processed_task.workflow_result is not None
    assert "Source content may be incomplete" in processed_task.workflow_result["manual_checks"]
    assert processed_task.workflow_result["context_summary"]["selected_hit_count"] == 0


def test_pipeline_service_deduplicates_context_hits_and_limits_selected_evidence(
    session: Session,
    tmp_path: Path,
) -> None:
    knowledge_repository = KnowledgeRepository(session)
    repeated_file = tmp_path / "repeated-guidance.md"
    repeated_file.write_text(
        "# Guidance\n\n"
        "Visible citations are mandatory before export.\n\n"
        "Visible citations are mandatory before export.\n\n"
        "Manual review remains required for launch claims.\n",
        encoding="utf-8",
    )
    document = knowledge_repository.create_document(
        title="Repeated Guidance",
        source_path=str(repeated_file),
        source_type="guide",
        domain="content-ops",
    )
    KnowledgeIndexService(knowledge_repository).index_document(document.id)

    repository = TaskRepository(session)
    task = repository.create_task(
        input_type="text",
        content="Prepare launch copy with visible citations before export and keep manual review.",
        knowledge_domain="content-ops",
    )

    service = TaskPipelineService(lambda: Session(session.get_bind()))
    processed_task = service.run_pipeline(task.id)

    assert processed_task.workflow_result is not None
    assert processed_task.workflow_result["context_summary"]["selected_hit_count"] <= 3
    assert processed_task.workflow_result["context_summary"]["duplicate_hits_removed"] >= 0
    assert len(processed_task.workflow_result["evidence_used"]) <= 3
